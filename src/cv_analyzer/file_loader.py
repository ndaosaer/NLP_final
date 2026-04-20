"""
Module de chargement de fichiers CV
Supporte les formats : PDF, Word (.docx), TXT
"""

import os
from pathlib import Path
from typing import Optional


class FileLoader:
    """
    Classe pour charger des CV depuis différents formats de fichiers.
    Formats supportés : PDF, DOCX, TXT
    """

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt'}

    def __init__(self):
        """Initialise le FileLoader et vérifie les dépendances."""
        self._check_dependencies()

    def _check_dependencies(self) -> None:
        """Vérifie que les bibliothèques nécessaires sont installées."""
        self._has_pdfplumber = False
        self._has_docx = False

        try:
            import pdfplumber
            self._has_pdfplumber = True
        except ImportError:
            print("Warning: pdfplumber non installé. Support PDF désactivé.")
            print("Installer avec: poetry add pdfplumber")

        try:
            import docx
            self._has_docx = True
        except ImportError:
            print("Warning: python-docx non installé. Support Word désactivé.")
            print("Installer avec: poetry add python-docx")

    def load(self, file_path: str) -> Optional[str]:
        """
        Charge un fichier CV et retourne son contenu texte.

        Args:
            file_path: Chemin vers le fichier CV

        Returns:
            Contenu texte du CV ou None si erreur
        """
        path = Path(file_path)

        # Vérification de l'existence du fichier
        if not path.exists():
            raise FileNotFoundError(f"Fichier non trouvé: {file_path}")

        # Détection du format
        extension = path.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Format non supporté: {extension}. "
                f"Formats acceptés: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )

        # Chargement selon le format
        if extension == '.pdf':
            return self._load_pdf(path)
        elif extension == '.docx':
            return self._load_docx(path)
        elif extension == '.txt':
            return self._load_txt(path)

    def _load_pdf(self, path: Path) -> str:
        """
        Extrait le texte d'un fichier PDF.

        Args:
            path: Chemin vers le fichier PDF

        Returns:
            Texte extrait du PDF
        """
        if not self._has_pdfplumber:
            raise ImportError(
                "pdfplumber requis pour les PDF. "
                "Installer avec: poetry add pdfplumber"
            )

        import pdfplumber

        text_content = []

        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)

        return '\n'.join(text_content)

    def _load_docx(self, path: Path) -> str:
        """
        Extrait le texte d'un fichier Word (.docx).

        Args:
            path: Chemin vers le fichier DOCX

        Returns:
            Texte extrait du document Word
        """
        if not self._has_docx:
            raise ImportError(
                "python-docx requis pour les fichiers Word. "
                "Installer avec: poetry add python-docx"
            )

        import docx

        document = docx.Document(path)

        text_content = []

        # Extraction des paragraphes
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        # Extraction des tableaux (souvent utilisés dans les CV)
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(' | '.join(row_text))

        return '\n'.join(text_content)

    def _load_txt(self, path: Path) -> str:
        """
        Charge le contenu d'un fichier texte.

        Args:
            path: Chemin vers le fichier TXT

        Returns:
            Contenu du fichier texte
        """
        # Essayer plusieurs encodages courants
        encodings = ['utf-8', 'latin-1', 'cp1252']

        for encoding in encodings:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue

        raise ValueError(f"Impossible de décoder le fichier {path}")

    def detect_format(self, file_path: str) -> str:
        """
        Détecte le format d'un fichier.

        Args:
            file_path: Chemin vers le fichier

        Returns:
            Extension du fichier (ex: '.pdf')
        """
        return Path(file_path).suffix.lower()

    def is_supported(self, file_path: str) -> bool:
        """
        Vérifie si un format de fichier est supporté.

        Args:
            file_path: Chemin vers le fichier

        Returns:
            True si le format est supporté
        """
        extension = self.detect_format(file_path)
        return extension in self.SUPPORTED_EXTENSIONS


def load_cv(file_path: str) -> str:
    """
    Fonction utilitaire pour charger un CV.

    Args:
        file_path: Chemin vers le fichier CV

    Returns:
        Contenu texte du CV
    """
    loader = FileLoader()
    return loader.load(file_path)


# Pour tester le module directement
if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python file_loader.py <chemin_fichier>")
        print("Formats supportés: PDF, DOCX, TXT")
        sys.exit(1)

    file_path = sys.argv[1]

    try:
        loader = FileLoader()
        content = loader.load(file_path)
        print(f"=== Contenu de {file_path} ===")
        print(content[:2000] if len(content) > 2000 else content)
        print(f"\n=== Longueur totale: {len(content)} caractères ===")
    except Exception as e:
        print(f"Erreur: {e}")
